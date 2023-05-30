from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from datetime import datetime


def style_ranking_numbers(rank: str) -> str:
    """
    Styles the rank number text

    Expects: Strings containing numbers
    Returns: 0s as "--" and all other numbers the same
    """

    if rank != '0':
        return rank
    
    return '--'

def set_cell_font_size(cell: type, font_size: int):
    """
    Sets cell paragraph's font size.

    Expects: Reference to Cell type, Integer > 0
    Returns: Nothing
    """

    text = cell.paragraphs[0]
    text.style.font.size = Pt(font_size)

def center_cell_text(cell: type) -> None:
    """
    Applies center styling to inner-cell text.

    Expects: Reference to Cell type
    Returns: Nothing
    """

    text = cell.paragraphs[0]
    text.alignment = WD_TABLE_ALIGNMENT.CENTER

def bolden_cell_text(cell:type) -> None:
    """
    Bolden the cell's paragraph text.

    Expects: Reference to Cell type
    Returns: Nothing
    """

    cell.paragraphs[0].runs[0].font.bold = True

def style_publish_date(date_string:str) -> str:

    new_date = datetime.strptime(date_string, '%Y-%m-%d')
    publish_date = new_date.strftime('%B %d %Y')

    return publish_date

def generate_namedate_header(unformatted_date_string: str, name_string: str) -> str:
    """
    Formats the information for use as a table header.

    Expects: Date as string, name as string
    Returns: Strings with format DD/DD/YYY ~newline~ LIST TITLE
    """
     
    date_string = style_publish_date(unformatted_date_string)
    return f"{date_string}\n{name_string}"