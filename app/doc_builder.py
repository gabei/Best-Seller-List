from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import stylizer


class DocBuilder:
    """This class is used to organize methods related to .docx document creation. Upon initialization, it creates its own document using the docx module."""

    def __init__(self):
        self.__doc = Document()

    def define_margins_inches(self, margin_list: list) -> list:
        """
        Sets the document margins in clock-wise order from margin_list, starting from left and ending on bottom.

        Expects: A list of four page margins in inches.
        Returns: The margin list if successful.
        """
        
        if len(margin_list) != 4:
            raise ValueError("List of margins must include 4 values.")

        sections = self.__doc.sections
        for section in sections:
            section.left_margin = Inches(margin_list[0])
            section.top_margin = Inches(margin_list[1])
            section.right_margin = Inches(margin_list[2])
            section.bottom_margin = Inches(margin_list[3])

        return margin_list

    def set_document_default_font_and_size(self, font_name="Arial", font_size="10"):
        """
        Sets the default font and font size for the document.

        Expects: Font name as a string, font size as an int.
        Returns: Tuple containing both values if successful.
        """

        style = self.__doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        return (font_name, font_size)

    def create_header(self, text):
        """Builds a document header using the passed text.

        - Expects a string
        - Returns the text if successful.
        """

        if type(text) is not str:
            raise TypeError("This method only accepts strings.")

        self.title = text

        section = self.__doc.sections[0]
        header = section.header
        stylizer.center_cell_text(header)
        stylizer.set_cell_font_size(header, 18)
        paragraph = header.paragraphs[0]
        paragraph.text = text
        

        #self.__doc.add_heading(text, 1)
        #TODO Heading style
        '''
        title_style = heading.style
        rFonts = title_style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Arial")
        '''
        
        return text
    
    def create_table(self, r: int, c: int):
        """
        Creates a table of specified rows and columns. Creates a blank table if not specified.

        - Expects integers > 0 for rows and cols
        - Invalid input will create a blank table
        - Returns 1 if successful
        """

        self.table = self.__doc.add_table(rows=r, cols=c)
        return 1

    def create_headers_for_table(self, header_list: list) -> None:
        """
        Sets the text of the header cells of the table in header_list order: left -> right

        Expects: A list of strings
        Returns: The list if successful
        """

        # Define the header cells of the table
        headers = self.table.rows[0].cells

        # Check for erroneous values
        if len(header_list) <= 0 or type(header_list) is not list:
            for index in range(0, len(headers)-1):
                headers[index].text = str(index + 1)
        
        # Apply List to Headers
        for cell, value in enumerate(header_list):
            table_cell = headers[cell]
            table_cell.text = value
            stylizer.center_cell_text(table_cell)
            stylizer.bolden_cell_text(table_cell)

        return header_list
    
    def set_column_widths_inches(self, width_list: list) -> list:
        """
        Sets table column widths in order of those supplied by width_list.

        Expects: List of widths in INCHES
        Returns: The list of lengths if successful.
        """

        cols = self.table.columns

        if not self.table: 
            raise ValueError('No table exists.')
        
        if len(width_list) != len(cols):
            raise ValueError('Number of columns supplied does not match table width.')
        
        # Set column width to all cells in each individual column
        for idx in range(0, len(width_list)):
            for cell_idx in range(0, len(cols[idx].cells)):
                cols[idx].cells[cell_idx].width = Inches(width_list[idx])
        
        return width_list
        
    def build_row_from(self, data):
        """
        Builds a row of table data using the book info passed.

        Expects an object
            - NYT Best Seller Data:
            - Rank, title, author, description, Rank last week, Weeks on list
        Returns said object if successful
        """

        row = self.table.add_row()
        
        row.cells[0].text = str(data['rank'])
        row.cells[1].text = f"{data['title']}, by {data['author']}. {data['description']}"
        row.cells[2].text = stylizer.style_ranking_numbers(str(data['rank_last_week']))
        row.cells[3].text = stylizer.style_ranking_numbers(str(data['weeks_on_list']))

        #TODO do not hardcode this???
        stylizer.center_cell_text(row.cells[0])
        stylizer.center_cell_text(row.cells[2])
        stylizer.center_cell_text(row.cells[3])

        return data

    def save_document(self, file_name: str) -> str:
        self.__doc.save(file_name+".docx")
        return f"Document saved at \"{file_name}.docx\""

    def __str__(self):
        return f"Document titled \"{self.title}\""
