from docx import Document


doc = Document()

# Heading 
doc.add_heading("The New York Times Best Seller List")
doc.add_page_break()

# Table
doc.add_table(rows=16, cols=4)
