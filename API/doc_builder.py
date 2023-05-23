from docx import Document
from docx.shared import Inches


class DocBuilder:
    """This class is used to organize methods related to .docx document creation. Upon initialization, it creates its own document using the docx module."""

    def __init__(self):
        self.doc = Document()

    def create_header(self, text):
        """Builds a document header using the passed text.
        - Expects a string
        - Returns the text if successful.
        """

        if type(text) is not str:
            raise TypeError("This method only accepts strings.")

        self.doc.add_heading(text)
        self.title = text
        return text
    
    def define_margins_inches(self, margin_list: list) -> list:
        """
        Sets the document margins in clock-wise order from margin_list, starting from left and ending on bottom.
        Expects: A list of four page margins in inches.
        Returns: The margin list if successful.
        """
        
        if len(margin_list) != 4:
            raise ValueError("List of margins must include 4 values.")

        sections = self.doc.sections
        for section in sections:
            section.left_margin = Inches(margin_list[0])
            section.top_margin = Inches(margin_list[0])
            section.right_margin = Inches(margin_list[0])
            section.bottom_margin = Inches(margin_list[0])

        return margin_list

    def create_table(self, r: int, c: int):
        """
        Creates a table of specified rows and columns. Creates a blank table if not specified.
        - Expects integers > 0 for rows and cols
        - Invalid input will create a blank table
        - Returns 1 if successful
        """
        self.table = self.doc.add_table(rows=r, cols=c)

        return 1

    def create_headers_for_table(self, header_list: list) -> None:
        # Define the headers in the table
        headers = self.table.rows[0].cells

        # Check for erroneous values
        if len(header_list) <= 0 or type(header_list) is not list:
            for index in range(0, len(headers)-1):
                headers[index].text = str(index + 1)

        # Fill in the headers if all is OK
        #for index in range(0, len(header_list)):
            #headers[index].text = header_list[index]

        for index, value in enumerate(header_list):
            headers[index].text = value

        return 1
    
    def set_column_widths_inches(self, width_list: list) -> list:
        """
        Sets table column widths in order of those supplied by width_list.
        Expects: List of widths in INCHES
        Returns: The list of lengths if successful.
        """
        cols = self.table.columns

        if not self.table: 
            raise ValueError('No table exists.')
        
        if len(width_list) != len(cols):
            raise ValueError('Number of columns supplied does not match table width.')
        
        # Set column width to all cells in each individual column
        for idx in range(0, len(width_list)):
            for cell_idx in range(0, len(cols[idx].cells)):
                cols[idx].cells[cell_idx].width = Inches(width_list[idx])
        
        return width_list
        

    def build_row_from(self, data):
        """Builds a row of table data using the data passed.
        - Expects an array in desired order"""

        row = self.table.add_row()
        # Rank, title-author-publisher-description, Rank last week, Weeks on list
        #print(data)
        
        row.cells[0].text = str(data['rank'])
        row.cells[1].text = f"{data['title']}, by {data['author']}. ({data['publisher']}.) {data['description']}"
        row.cells[2].text = str(data['rank_last_week'])
        row.cells[3].text = str(data['weeks_on_list'])

        return data

    def save_document(self, file_name: str) -> str:
        self.doc.save(file_name+".docx")
        return f"Document saved at \"{file_name}.docx\""

    def __str__(self):
        return f"Document titled \"{self.title}\""
